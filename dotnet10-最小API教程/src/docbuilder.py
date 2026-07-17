# -*- coding: utf-8 -*-
"""
docbuilder.py —— 《ASP.NET Core 最小 API 完全教程（.NET 10）》Word 生成工具库
提供：Word 排版辅助函数 + matplotlib 中文示意图绘制辅助。
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager as fm
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle

# ---------------------------------------------------------------------------
# 路径与字体
# ---------------------------------------------------------------------------
BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE, "images")
os.makedirs(IMG_DIR, exist_ok=True)

FONT_PATH = "/projects/sandbox/fonts/NotoSansSC.ttf"
fm.fontManager.addfont(FONT_PATH)
_CN_PROP = fm.FontProperties(fname=FONT_PATH)
CN_NAME = _CN_PROP.get_name()
plt.rcParams["font.family"] = CN_NAME
plt.rcParams["axes.unicode_minus"] = False

# 文档中使用的字体
BODY_CN = "Microsoft YaHei"     # 正文中文
BODY_EN = "Segoe UI"            # 正文英文
CODE_EN = "Consolas"            # 代码英文
CODE_CN = "Microsoft YaHei"     # 代码中文注释

# 配色
CLR_TITLE   = RGBColor(0x1F, 0x3A, 0x5F)
CLR_H1      = RGBColor(0x0E, 0x4C, 0x92)
CLR_H2      = RGBColor(0x1B, 0x6E, 0xC2)
CLR_H3      = RGBColor(0x2E, 0x74, 0xB5)
CLR_BODY    = RGBColor(0x22, 0x22, 0x22)
CLR_CODE    = RGBColor(0x1E, 0x1E, 0x1E)
CLR_NOTE    = RGBColor(0x6A, 0x4A, 0x00)
CLR_CAPTION = RGBColor(0x70, 0x70, 0x70)

# 图表配色
DZ = dict(
    blue="#DCEBFA", blue_e="#2E74B5",
    green="#DDF3E4", green_e="#2E9E5B",
    orange="#FDECD2", orange_e="#E08E0B",
    gray="#EFEFEF", gray_e="#8C8C8C",
    purple="#EADCF7", purple_e="#8E44AD",
    red="#FADBD8", red_e="#C0392B",
    yellow="#FCF3CF", yellow_e="#C9A227",
)


# ===========================================================================
# Word 排版辅助
# ===========================================================================
def _set_run(run, ascii_font, cjk_font, size, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), cjk_font)


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _border(paragraph, color="CCCCCC", size="6", sides=("top", "bottom", "left", "right")):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for s in sides:
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), color)
        pbdr.append(e)
    pPr.append(pbdr)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    if level == 1:
        p.paragraph_format.space_before = Pt(22)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        _set_run(r, BODY_EN, BODY_CN, 20, bold=True, color=CLR_H1)
        # 底部横线
        _border(p, color="0E4C92", size="12", sides=("bottom",))
    elif level == 2:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        _set_run(r, BODY_EN, BODY_CN, 15.5, bold=True, color=CLR_H2)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        _set_run(r, BODY_EN, BODY_CN, 13, bold=True, color=CLR_H3)
    return p


def para(doc, runs, size=11, align=None, space_after=8, first_indent=False):
    """runs: str 或 [(text, {bold,italic,color,mono})]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opt in runs:
        r = p.add_run(text)
        if opt.get("mono"):
            _set_run(r, CODE_EN, CODE_CN, size - 0.5,
                     bold=opt.get("bold", False), color=opt.get("color", CLR_CODE))
            _shade_run(r, "F0F0F0")
        else:
            _set_run(r, BODY_EN, BODY_CN, size,
                     bold=opt.get("bold", False),
                     italic=opt.get("italic", False),
                     color=opt.get("color", CLR_BODY))
    return p


def _shade_run(run, fill):
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def bullet(doc, runs, level=0, size=11):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opt in runs:
        r = p.add_run(text)
        if opt.get("mono"):
            _set_run(r, CODE_EN, CODE_CN, size - 0.5, bold=opt.get("bold", False),
                     color=opt.get("color", CLR_CODE))
            _shade_run(r, "F0F0F0")
        else:
            _set_run(r, BODY_EN, BODY_CN, size, bold=opt.get("bold", False),
                     italic=opt.get("italic", False), color=opt.get("color", CLR_BODY))
    return p


def numbered(doc, runs, size=11):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opt in runs:
        r = p.add_run(text)
        _set_run(r, BODY_EN, BODY_CN, size, bold=opt.get("bold", False),
                 italic=opt.get("italic", False), color=opt.get("color", CLR_BODY))
    return p


def code_block(doc, code, lang="csharp"):
    """整段代码块：浅灰底 + 边框 + 等宽字体。"""
    lines = code.split("\n")
    # 顶部语言标签
    lab = doc.add_paragraph()
    lab.paragraph_format.space_before = Pt(6)
    lab.paragraph_format.space_after = Pt(0)
    lr = lab.add_run(f"  {lang}")
    _set_run(lr, CODE_EN, CODE_CN, 8.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    _shade(lab, "2E74B5")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    _shade(p, "F4F5F7")
    _border(p, color="C9CDD4", size="6")
    for i, line in enumerate(lines):
        r = p.add_run(line)
        _set_run(r, CODE_EN, CODE_CN, 9.5, color=CLR_CODE)
        if i < len(lines) - 1:
            r.add_break()
    return p


def inline_code(text):
    return (text, {"mono": True})


def note(doc, title, text, kind="tip"):
    palette = {
        "tip":  ("E8F4FD", "2E74B5", "提示"),
        "warn": ("FDECEA", "C0392B", "注意"),
        "key":  ("FFF7E0", "C9A227", "重点"),
        "info": ("EDEDED", "666666", "说明"),
    }
    fill, bar, deflabel = palette.get(kind, palette["tip"])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    _shade(p, fill)
    _border(p, color=bar, size="6")
    r = p.add_run(f"【{title or deflabel}】 ")
    _set_run(r, BODY_EN, BODY_CN, 10.5, bold=True, color=RGBColor.from_string(bar))
    r2 = p.add_run(text)
    _set_run(r2, BODY_EN, BODY_CN, 10.5, color=CLR_BODY)
    return p


def image(doc, path, caption=None, width=6.0):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(12)
        r = c.add_run(caption)
        _set_run(r, BODY_EN, BODY_CN, 9.5, italic=True, color=CLR_CAPTION)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        para_cell = hdr[i].paragraphs[0]
        r = para_cell.add_run(h)
        _set_run(r, BODY_EN, BODY_CN, 10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade(para_cell, "2E74B5")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pc = cells[i].paragraphs[0]
            r = pc.add_run(val)
            _set_run(r, BODY_EN, BODY_CN, 10, color=CLR_BODY)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    return t


def page_break(doc):
    doc.add_page_break()


# ===========================================================================
# matplotlib 示意图辅助
# ===========================================================================
def _box(ax, x, y, w, h, text, fc, ec, fs=12, bold=False, text_color="#1a1a1a",
         rounding=0.02, lw=1.8, align="center"):
    box = FancyBboxPatch((x, y), w, h,
                         boxstyle=f"round,pad=0.01,rounding_size={rounding}",
                         linewidth=lw, edgecolor=ec, facecolor=fc, zorder=2)
    ax.add_patch(box)
    ha = {"center": "center", "left": "left"}[align]
    tx = x + w / 2 if align == "center" else x + 0.02
    ax.text(tx, y + h / 2, text, ha=ha, va="center", fontsize=fs,
            color=text_color, fontweight="bold" if bold else "normal", zorder=3,
            linespacing=1.4)
    return box


def _arrow(ax, x1, y1, x2, y2, color="#444", text=None, text_off=(0, 0.03),
           style="-|>", lw=2.0, fs=10, ls="-"):
    ar = FancyArrowPatch((x1, y1), (x2, y2), arrowstyle=style,
                         mutation_scale=18, color=color, lw=lw,
                         linestyle=ls, zorder=1)
    ax.add_patch(ar)
    if text:
        mx, my = (x1 + x2) / 2 + text_off[0], (y1 + y2) / 2 + text_off[1]
        ax.text(mx, my, text, ha="center", va="center", fontsize=fs, color=color)


def _new_ax(w=11, h=6):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10 * h / w)
    ax.axis("off")
    return fig, ax


def save(fig, name):
    path = os.path.join(IMG_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path
