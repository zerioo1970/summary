# -*- coding: utf-8 -*-
"""定制 pandoc 参考样式模板 reference.docx：中文字体、代码块底色、标题配色、提示框、A4页面、页脚。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REF = "/projects/sandbox/summary/dotnet10-最小API教程/template/reference.docx"
CJK = "Microsoft YaHei"
EN = "Segoe UI"
MONO = "Consolas"

doc = Document(REF)


def set_font(style, ascii_font, cjk_font, size=None, bold=None, color=None, italic=None):
    f = style.font
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = color
    f.name = ascii_font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), cjk_font)
    rfonts.set(qn("w:cs"), ascii_font)


def shade(style, fill):
    ppr = style.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def borders(style, color, sz="6", sides=("top", "bottom", "left", "right"), space="8"):
    ppr = style.element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for s in sides:
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), space); e.set(qn("w:color"), color)
        pbdr.append(e)
    ppr.append(pbdr)


def spacing(style, before=None, after=None, line=None):
    pf = style.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


styles = doc.styles

# 正文
set_font(styles["Normal"], EN, CJK, size=11, color=RGBColor(0x22, 0x22, 0x22))
spacing(styles["Normal"], after=6, line=1.3)

# 标题配色与字号
if "Title" in [s.name for s in styles]:
    set_font(styles["Title"], EN, CJK, size=30, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))
for name, size, color in [
    ("Heading 1", 20, RGBColor(0x0E, 0x4C, 0x92)),
    ("Heading 2", 15.5, RGBColor(0x1B, 0x6E, 0xC2)),
    ("Heading 3", 13, RGBColor(0x2E, 0x74, 0xB5)),
    ("Heading 4", 11.5, RGBColor(0x2E, 0x74, 0xB5)),
]:
    try:
        set_font(styles[name], EN, CJK, size=size, bold=True, color=color)
        spacing(styles[name], before=14 if "1" in name else 10, after=6)
    except KeyError:
        pass

# 代码块：Source Code 段落样式（浅灰底 + 边框 + 等宽）
for scname in ["Source Code", "SourceCode", "Verbatim"]:
    try:
        sc = styles[scname]
        set_font(sc, MONO, CJK, size=9.5, color=RGBColor(0x1E, 0x1E, 0x1E))
        shade(sc, "F4F5F7")
        borders(sc, "C9CDD4", sz="4", space="2")
        spacing(sc, before=0, after=0, line=1.15)
    except KeyError:
        pass

# 行内代码 Verbatim Char
for vcname in ["Verbatim Char", "VerbatimChar"]:
    try:
        set_font(styles[vcname], MONO, CJK, size=10, color=RGBColor(0xC7, 0x25, 0x4E))
    except KeyError:
        pass

# 引用块（提示框）：Block Text / Quote
for qname in ["Block Text", "BlockText", "Quote"]:
    try:
        q = styles[qname]
        set_font(q, EN, CJK, size=10.5, color=RGBColor(0x33, 0x33, 0x33))
        shade(q, "EAF3FB")
        borders(q, "2E74B5", sz="18", sides=("left",), space="10")
        spacing(q, before=6, after=6, line=1.3)
    except KeyError:
        pass

# 图片说明（Caption）居中灰色斜体
try:
    cap = styles["Image Caption"]
    set_font(cap, EN, CJK, size=9.5, italic=True, color=RGBColor(0x70, 0x70, 0x70))
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
except KeyError:
    pass

# 页面 A4 + 页边距 + 页脚
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.left_margin = sec.right_margin = Inches(0.9)
sec.top_margin = sec.bottom_margin = Inches(0.9)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("《ASP.NET Core 最小 API 完全教程（.NET 10）》")
r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
r.font.name = EN
rpr = r._element.get_or_add_rPr()
rf = OxmlElement("w:rFonts"); rf.set(qn("w:eastAsia"), CJK); rpr.append(rf)

doc.save(REF)
print("参考模板已定制:", REF)
