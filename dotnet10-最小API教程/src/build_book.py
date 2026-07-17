# -*- coding: utf-8 -*-
"""《ASP.NET Core 最小 API 完全教程（.NET 10）》完整文档生成总脚本。
每次运行重生成整本 docx，保证结构一致。当前已写：封面、目录、第1~3章、附录A。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import docbuilder as db
from docbuilder import (heading, para, bullet, numbered, code_block, note,
                        image, table, page_break, _set_run,
                        BODY_CN, BODY_EN, CLR_TITLE, CLR_H1, CLR_H2, CLR_CAPTION, IMG_DIR)

C = lambda t: (t, {})
B = lambda t: (t, {"bold": True})
M = lambda t: (t, {"mono": True})


def img(doc, name, caption, width=6.0):
    image(doc, os.path.join(IMG_DIR, name), caption, width)


def add_footer(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《ASP.NET Core 最小 API 完全教程（.NET 10）》")
    _set_run(r, BODY_EN, BODY_CN, 8.5, color=CLR_CAPTION)


# ===========================================================================
def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    for txt, sz in [("ASP.NET Core", 30), ("最小 API 完全教程", 34)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); _set_run(r, BODY_EN, BODY_CN, sz, bold=True, color=CLR_TITLE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Minimal API · 基于 .NET 10"); _set_run(r, BODY_EN, BODY_CN, 18, color=CLR_H2)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("从零基础到进阶实战 · 图文并茂 · 大量可运行示例")
    _set_run(r, BODY_EN, BODY_CN, 13, italic=True, color=db.CLR_BODY)
    for _ in range(6):
        doc.add_paragraph()
    for line in ["面向读者：完全没接触过最小 API 的新手，以及想系统梳理的开发者",
                 "环境版本：.NET 10 SDK / C# 14",
                 "编写工具：Visual Studio 2026 或 VS Code 均可"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); _set_run(r, BODY_EN, BODY_CN, 11, color=CLR_CAPTION)
    page_break(doc)


def build_toc(doc):
    heading(doc, "目录", 1)
    toc = [
        ("第一部分  入门与原理", None),
        ("第 1 章  认识最小 API", ["1.1 什么是最小 API", "1.2 最小 API vs 传统 MVC 控制器",
            "1.3 为什么 .NET 10 推荐用最小 API", "1.4 .NET 10 带来的新变化", "1.5 适用场景与不适用场景"]),
        ("第 2 章  运行原理与请求处理管线", ["2.1 一个 HTTP 请求在 ASP.NET Core 中的旅程",
            "2.2 WebApplication 与 WebApplicationBuilder 的关系", "2.3 中间件（Middleware）管道原理",
            "2.4 主机、Kestrel、依赖注入容器三者关系", "2.5 Program.cs 顶级语句逐行拆解"]),
        ("第 3 章  环境搭建与第一个程序", ["3.1 安装 .NET 10 SDK", "3.2 用 CLI 创建项目 dotnet new web",
            "3.3 项目结构详解", "3.4 运行与热重载 dotnet watch", "3.5 Hello World 全流程示例"]),
        ("第二部分  核心功能", None),
        ("第 4 章  路由（Routing）", None),
        ("第 5 章  参数绑定（Parameter Binding）", None),
        ("第 6 章  返回结果（Results）", None),
        ("第 7 章  依赖注入（DI）", None),
        ("第三部分  数据与校验", None),
        ("第 8 章  请求校验（.NET 10 新特性重点）", None),
        ("第 9 章  数据持久化与 EF Core", None),
        ("第四部分  文档、安全与实时通信", None),
        ("第 10 章  OpenAPI 与接口文档（.NET 10 新特性）", None),
        ("第 11 章  认证与授权", None),
        ("第 12 章  端点过滤器与中间件", None),
        ("第 13 章  实时通信：Server-Sent Events（.NET 10 新特性）", None),
        ("第五部分  工程化与进阶", None),
        ("第 14 章  配置、日志与错误处理", None),
        ("第 15 章  跨域（CORS）、限流与健康检查", None),
        ("第 16 章  API 版本控制", None),
        ("第 17 章  测试", None),
        ("第 18 章  组织大型项目", None),
        ("第 19 章  部署与性能", None),
        ("第 20 章  综合实战：从零构建一个完整 API", None),
        ("附录", None),
        ("附录 A  前端与后端如何协作（含所有前端调用方法与示例）",
            ["A.1 前端、后端分别是什么", "A.2 最小 API 的定位：只提供数据", "A.3 HTTP 协议入门",
             "A.4 一次完整调用的全流程", "A.5 前后端分离 vs 前后端一体",
             "A.6 前端调用后端的各种方法大全（各举一例）", "A.7 没有前端时如何测试接口", "A.8 跨域（CORS）初识"]),
        ("附录 B  最小 API 常用速查表", None),
        ("附录 C  从 Controller 迁移到最小 API 对照表", None),
        ("附录 D  常见错误与排查", None),
    ]
    for title, subs in toc:
        is_part = ("部分" in title or title == "附录") and "章" not in title
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2 if subs else 3)
        r = p.add_run(title)
        if is_part:
            p.paragraph_format.space_before = Pt(10)
            _set_run(r, BODY_EN, BODY_CN, 13, bold=True, color=CLR_TITLE)
        else:
            _set_run(r, BODY_EN, BODY_CN, 11.5, bold=True, color=CLR_H1)
        if subs:
            for s in subs:
                sp = doc.add_paragraph(); sp.paragraph_format.left_indent = Pt(22)
                sp.paragraph_format.space_after = Pt(1)
                sr = sp.add_run(s); _set_run(sr, BODY_EN, BODY_CN, 10, color=db.CLR_BODY)
    note(doc, "阅读提示", "本教程前三章重在讲清“是什么、为什么、怎么跑起来”，从第 4 章开始进入具体功能。"
         "关于“前端如何调用这些接口、前后端怎么协作”的内容，集中放在书末【附录 A】，"
         "初学者可在读完第 3 章后先翻阅附录 A 建立整体认知，再回到第 4 章深入。"
         "标有“.NET 10 新特性”的章节是相较旧版本的重要升级点，建议重点掌握。", "info")
    page_break(doc)


# ===========================================================================
def build_chapter1(doc):
    heading(doc, "第 1 章  认识最小 API", 1)
    para(doc, [C("在写第一行代码之前，我们先花一章时间把两个最基本的问题弄明白："),
               B("最小 API 到底是什么？"), C("以及"), B("我为什么要用它？"),
               C("想清楚这两点，后面学起来才不会“知其然不知其所以然”。")])

    heading(doc, "1.1 什么是最小 API（Minimal API）", 2)
    para(doc, [C("“最小 API”是 ASP.NET Core 提供的一种"), B("用最少的代码构建 HTTP 接口"),
               C("的方式。它的英文是 "), M("Minimal API"),
               C("，从 .NET 6 引入，到 .NET 10 已经非常成熟，成为微软官方推荐的 API 构建方式。")])
    para(doc, [C("“最小”体现在：你不需要建一堆文件、写一堆类，"),
               B("一个 Program.cs 文件、几行代码"), C("，就能跑起来一个真正的 Web 接口。看下面这段完整程序：")])
    code_block(doc, '''var builder = WebApplication.CreateBuilder(args);
var app = builder.Build();

// 定义一个接口：访问 /hello 时返回一段文字
app.MapGet("/hello", () => "你好，最小 API！");

app.Run();  // 启动服务器，开始监听请求''')
    para(doc, [C("就这么几行，运行后在浏览器打开 "), M("http://localhost:5000/hello"),
               C("，你就能看到 “你好，最小 API！”。这就是一个完整的后端接口了——"),
               B("没有 Controller 类，没有一堆配置文件，没有繁琐的约定。")])
    note(doc, "名词解释", "“API”指应用程序编程接口，这里可以简单理解为“一个可以通过网址访问、"
         "用来交换数据的服务入口”。“HTTP API”就是通过 HTTP 协议访问的接口。", "info")

    heading(doc, "1.2 最小 API vs 传统 MVC 控制器", 2)
    para(doc, [C("在最小 API 出现之前，ASP.NET Core 构建接口的主流方式是"),
               B("MVC 控制器（Controller）"), C("。两者都能做出一样的接口，区别在于“写法的繁简”和“适用的规模”。")])
    img(doc, "fig_1_1_minimal_vs_mvc.png", "图 1-1  最小 API 与传统 MVC 控制器的对比")
    para(doc, [C("下面用一个“返回商品列表”的例子直观对比。先看"), B("传统 MVC 控制器"), C("的写法：")])
    code_block(doc, '''// 传统 MVC：需要单独的 Controller 类、特性标注、约定命名
[ApiController]
[Route("api/[controller]")]
public class ProductsController : ControllerBase
{
    [HttpGet]
    public IActionResult GetAll()
    {
        var products = new[] { "苹果", "香蕉" };
        return Ok(products);
    }
}''')
    para(doc, [C("再看"), B("最小 API"), C("实现同样的功能：")])
    code_block(doc, '''// 最小 API：一行搞定，无需类、无需特性
app.MapGet("/api/products", () => new[] { "苹果", "香蕉" });''')
    para(doc, [C("两者对外暴露的接口效果"), B("完全一样"),
               C("，但最小 API 明显更简洁。下表总结它们的核心差异：")])
    table(doc, ["对比项", "最小 API", "MVC 控制器"],
          [["代码量", "极少，几行即可", "较多，需类+特性+约定"],
           ["文件组织", "可集中在 Program.cs", "通常每个控制器一个文件"],
           ["学习曲线", "平缓，直观", "较陡，需理解约定"],
           ["启动/内存", "更快、更省", "相对更重"],
           ["适合规模", "中小型、微服务", "大型、复杂系统"],
           ["功能完整度", "已覆盖绝大多数场景", "最全面"]])
    note(doc, "重点", "最小 API 并不是“功能被阉割的 MVC”。到了 .NET 10，它已经补齐了校验、"
         "OpenAPI 文档、认证授权等企业级能力，绝大多数项目都可以放心使用。", "key")

    heading(doc, "1.3 为什么 .NET 10 推荐用最小 API", 2)
    para(doc, "微软在 .NET 10 官方文档中，明确把最小 API 列为构建 HTTP 接口的推荐方式。原因主要有四点：")
    numbered(doc, [B("上手快、心智负担低"), C("——新手不用先理解一堆 MVC 约定，看到路由就能对应到代码。")])
    numbered(doc, [B("性能好"), C("——更少的抽象层意味着更快的启动速度和更低的内存占用，天然适合云原生、容器与微服务。")])
    numbered(doc, [B("与现代实践契合"), C("——前后端分离、微服务架构下，后端往往只需提供 JSON 接口，最小 API 正好胜任。")])
    numbered(doc, [B("功能已足够成熟"), C("——经过 .NET 6/7/8/9 的持续增强，到 .NET 10 已是可用于生产的完整框架。")])

    heading(doc, "1.4 .NET 10 带来的新变化", 2)
    para(doc, "如果你之前用过旧版本的最小 API，或者看过一些旧教程，要特别注意 .NET 10 带来的几项重要升级：")
    bullet(doc, [B("内置模型校验"), C("：不用再自己写校验代码，直接用 "), M("[Required]"),
                 C("、"), M("[Range]"), C(" 等特性即可，且兼容原生 AoT 编译（第 8 章详讲）。")])
    bullet(doc, [B("OpenAPI 3.1 支持"), C("：内置生成符合 OpenAPI 3.1 规范的接口文档，能配合 Swagger UI / Scalar 使用（第 10 章）。")])
    bullet(doc, [B("Server-Sent Events（SSE）"), C("：通过 "), M("TypedResults.ServerSentEvents"),
                 C(" 轻松实现服务器实时推送（第 13 章）。")])
    bullet(doc, [B("原生 AoT 编译增强"), C("：可编译成独立的原生可执行文件，启动更快、体积更小（第 19 章）。")])
    note(doc, "注意", "由于最小 API 每个版本都在快速演进，网上不少旧教程的写法在 .NET 10 里可能已有更简单的替代方案。"
         "本教程所有示例均以 .NET 10 为准。", "warn")

    heading(doc, "1.5 适用场景与不适用场景", 2)
    para(doc, "任何技术都有它擅长和不擅长的地方。什么时候该用最小 API，什么时候该考虑别的方案？")
    img(doc, "fig_1_2_when_to_use.png", "图 1-2  最小 API 的适用与不适用场景")
    para(doc, [C("简单来说："), B("绝大多数“提供数据接口”的后端项目，最小 API 都是很好的选择"),
               C("。只有在需要服务器端渲染网页、或深度依赖 MVC 特有管线的少数场景，才需要另作考虑。"
                 "对新手而言，先用最小 API 把基础打牢，是最高效的路径。")])
    note(doc, "本章小结", "最小 API 是一种用极少代码构建 HTTP 接口的方式；它比 MVC 更轻更快，"
         "在 .NET 10 中功能已相当完整，是官方推荐的首选。记住一句话：最小 API 负责“提供数据接口”，"
         "而不是“渲染网页”。下一章，我们深入后端内部，看看一个请求进入最小 API 后到底经历了什么。"
         "（如果你想先搞清楚“前端是怎么调用这些接口的”，可以随时翻到书末的【附录 A】。）", "key")
    page_break(doc)


# ===========================================================================
def build_chapter2(doc):
    heading(doc, "第 2 章  运行原理与请求处理管线", 1)
    para(doc, [C("上一章我们知道了最小 API 能用几行代码建出接口。但这几行代码背后，程序究竟是怎么跑起来、"
                 "又是怎么处理一个请求的？这一章我们“掀开引擎盖”，把运行原理讲清楚。"),
               B("理解了原理，后面遇到任何问题都能心中有数、不再靠猜。")])

    heading(doc, "2.1 一个 HTTP 请求在 ASP.NET Core 中的旅程", 2)
    para(doc, [C("当浏览器向你的接口发出一个请求，这个请求并不是“直接”跳到你写的那行 "),
               M("MapGet"), C(" 代码里，而是要依次经过好几站，像坐地铁一样。看下图：")])
    img(doc, "fig_ch2_pipeline.png", "图 2-1  一个 HTTP 请求的完整旅程")
    para(doc, "我们顺着箭头，看请求经过的每一站：")
    numbered(doc, [B("客户端发起请求"), C("：浏览器 / App / Postman 向服务器地址发出 HTTP 请求。")])
    numbered(doc, [B("Kestrel Web 服务器"), C("：ASP.NET Core 内置的高性能 Web 服务器，它是请求进入程序的“大门”，"
                   "负责接收原始的网络数据并解析成 HTTP 请求对象。")])
    numbered(doc, [B("中间件管道"), C("：请求依次穿过一系列“中间件”，比如异常处理、HTTPS 重定向、认证授权等（下面 2.3 节详讲）。")])
    numbered(doc, [B("路由（Routing）"), C("：根据请求的方法和 URL（如 "), M("GET /api/products"),
                   C("），找到与之匹配的那个端点。")])
    numbered(doc, [B("端点处理函数"), C("：也就是你写在 "), M("MapGet(...)"),
                   C(" 里的那段代码，真正执行业务逻辑，生成结果。")])
    para(doc, [C("处理完后，"), B("响应会沿着原路“逆向”返回"),
               C("，再经过中间件、Kestrel，最终回到客户端。整个过程通常在几毫秒内完成。")])
    note(doc, "关键理解", "你写的代码（端点处理函数）只是这趟旅程的“最后一站”。它前面还有 Kestrel 和一整条中间件管道"
         "在默默工作。这就是为什么理解管线很重要。", "key")

    heading(doc, "2.2 WebApplication 与 WebApplicationBuilder 的关系", 2)
    para(doc, [C("回看第 1 章的代码，开头两行是："), ])
    code_block(doc, '''var builder = WebApplication.CreateBuilder(args);  // 第 1 行
var app = builder.Build();                          // 第 2 行''')
    para(doc, [C("这两行代表了程序的"), B("两个阶段"), C("：先用 "), M("builder"),
               C("（建造者）做各种准备工作，再调用 "), M("builder.Build()"), C(" 造出正式的应用对象 "),
               M("app"), C("。可以类比成“先备料、再开张营业”：")])
    img(doc, "fig_ch2_builder_app.png", "图 2-2  builder 与 app：配置阶段 → 运行阶段")
    heading(doc, "① 配置阶段：WebApplicationBuilder", 3)
    para(doc, [C("这个阶段你通过 "), M("builder"), C(" 做三件典型的事——它们都必须在 "),
               M("Build()"), C(" 之前完成：")])
    bullet(doc, [B("注册服务"), C("："), M("builder.Services.AddXxx()"), C("，把要用的功能（数据库、CORS、认证等）登记进依赖注入容器。")])
    bullet(doc, [B("读取配置"), C("："), M("builder.Configuration"), C("，读取 appsettings.json、环境变量等。")])
    bullet(doc, [B("设置日志"), C("："), M("builder.Logging"), C("，配置日志输出方式。")])
    heading(doc, "② 运行阶段：WebApplication", 3)
    para(doc, [C("调用 "), M("Build()"), C(" 后得到 "), M("app"),
               C("，这个阶段你做两件事，最后启动服务器：")])
    bullet(doc, [B("配置中间件"), C("："), M("app.UseXxx()"), C("，把中间件按顺序装进管道。")])
    bullet(doc, [B("映射路由"), C("："), M("app.MapGet / MapPost ..."), C("，定义具体的接口端点。")])
    bullet(doc, [B("启动监听"), C("："), M("app.Run()"), C("，程序开始监听端口、处理请求（这行之后的代码在服务器关闭前不会执行）。")])
    note(doc, "常见错误", "服务的注册（AddXxx）必须写在 builder.Build() 之前，中间件与路由（UseXxx / MapXxx）"
         "必须写在 Build() 之后。顺序搞反会直接报错。", "warn")

    heading(doc, "2.3 中间件（Middleware）管道原理", 2)
    para(doc, [C("中间件是 ASP.NET Core 最核心的概念之一。一个中间件就是一段“能处理请求的代码”，"
                 "多个中间件"), B("首尾相连组成一条管道"),
               C("。请求像水一样从管道一头流入，经过每一节，再从另一头流出。经典的示意是“洋葱模型”：")])
    img(doc, "fig_ch2_middleware.png", "图 2-3  中间件管道（洋葱模型）", width=4.6)
    para(doc, [C("请求"), B("由外向内"), C("逐层穿过每个中间件，到达最里面的端点（你的代码）；"
               "生成响应后，再"), B("由内向外"), C("逐层返回。每一层都可以：①对请求做点事；"
               "②决定是否把请求交给下一层；③在响应返回时再做点事。")])
    para(doc, "下面是一段典型的中间件配置，注意——它们的书写顺序，就是请求穿过的顺序：")
    code_block(doc, '''var app = builder.Build();

app.UseExceptionHandler("/error");  // 最外层：捕获后面所有环节的异常
app.UseHttpsRedirection();          // 把 http 请求重定向到 https
app.UseAuthentication();            // 认证：你是谁
app.UseAuthorization();             // 授权：你能不能访问

app.MapGet("/api/products", () => new[] { "苹果", "香蕉" });  // 端点

app.Run();''')
    para(doc, [C("你也可以用 "), M("app.Use(...)"), C(" 写一个自定义中间件。参数里的 "),
               M("next"), C(" 代表“管道的下一节”，调用它请求才会继续往里走：")])
    code_block(doc, '''app.Use(async (context, next) =>
{
    // ① 请求进入时（由外向内）
    Console.WriteLine($"收到请求：{context.Request.Path}");

    await next();  // 把请求交给下一节中间件 / 端点

    // ② 响应返回时（由内向外）
    Console.WriteLine($"返回状态码：{context.Response.StatusCode}");
});''')
    note(doc, "重点", "中间件的顺序至关重要。例如 UseAuthentication 必须在 UseAuthorization 之前，"
         "否则“还没搞清你是谁，就先判断你能不能进”，逻辑就错了。第 12 章会深入讲中间件与端点过滤器。", "key")

    heading(doc, "2.4 主机、Kestrel、依赖注入容器三者关系", 2)
    para(doc, [C("初学时经常听到几个词：主机（Host）、Kestrel、依赖注入容器（DI）。它们到底是什么关系？"
                 "一张图看懂：")])
    img(doc, "fig_ch2_host.png", "图 2-4  主机、Kestrel、DI 容器等组件的关系")
    bullet(doc, [B("主机（Host）"), C("：最外层的“大管家”，负责统一启动、托管和优雅关闭下面的所有组件。"
                 "你调用的 "), M("app.Run()"), C(" 本质就是启动主机。")])
    bullet(doc, [B("Kestrel"), C("：内置的跨平台 Web 服务器，真正负责监听端口、收发 HTTP 数据。")])
    bullet(doc, [B("依赖注入容器（DI Container）"), C("：管理程序里所有“服务”的仓库，需要用哪个服务，它就自动“送货上门”（第 7 章详讲）。")])
    bullet(doc, [B("配置系统与日志系统"), C("：分别负责读取参数配置、输出运行日志。")])
    para(doc, "简单说：主机把这些组件“装在一起、统一管理”，你写的最小 API 就运行在这套地基之上。")

    heading(doc, "2.5 Program.cs 顶级语句逐行拆解", 2)
    para(doc, [C("你可能注意到，最小 API 的 "), M("Program.cs"), C(" 里"),
               B("没有 class，也没有 Main 方法"),
               C("，代码直接“裸写”在文件里。这用到了 C# 的“顶级语句（Top-level statements）”特性——"
                 "编译器会自动帮你把这些代码包进一个隐藏的 Main 方法里。下面把一个稍完整的 Program.cs 逐行拆开讲：")])
    code_block(doc, '''// ① 创建建造者，args 是命令行参数
var builder = WebApplication.CreateBuilder(args);

// ② 【配置阶段】注册服务：这里注册了 CORS 支持
builder.Services.AddCors();

// ③ 造出应用对象
var app = builder.Build();

// ④ 【运行阶段】配置中间件
app.UseCors();

// ⑤ 映射一个 GET 端点，返回当前时间（自动序列化为 JSON）
app.MapGet("/api/time", () => new { now = DateTime.Now });

// ⑥ 启动服务器，开始监听（阻塞在此，直到程序关闭）
app.Run();''')
    numbered(doc, [M("CreateBuilder(args)"), C("：创建建造者，读取默认配置、日志等。")])
    numbered(doc, [M("AddCors()"), C("：把 CORS 服务登记到 DI 容器（属于配置阶段）。")])
    numbered(doc, [M("Build()"), C("：结束配置，生成 app。")])
    numbered(doc, [M("UseCors()"), C("：启用 CORS 中间件（属于运行阶段）。")])
    numbered(doc, [M("MapGet(...)"), C("：定义接口。返回的匿名对象会被自动转成 JSON。")])
    numbered(doc, [M("app.Run()"), C("：启动并阻塞，服务器开始对外服务。")])
    note(doc, "本章小结", "一个请求要经过 Kestrel → 中间件管道 → 路由 → 端点，再原路返回；"
         "程序分“配置阶段（builder）”和“运行阶段（app）”两步；中间件像洋葱一样层层包裹，顺序很关键；"
         "主机作为大管家托管 Kestrel、DI、配置、日志。搞懂这套原理，下一章我们就动手把开发环境搭起来，"
         "亲手跑通第一个程序。", "key")
    page_break(doc)


# ===========================================================================
def build_chapter3(doc):
    heading(doc, "第 3 章  环境搭建与第一个程序", 1)
    para(doc, [C("原理讲完，该动手了。这一章带你从零把环境搭好，并"), B("亲手跑通第一个最小 API 程序"),
               C("。所有命令都基于跨平台通用的 "), M(".NET CLI"),
               C("（命令行工具），无论你用 Visual Studio 2026 还是 VS Code 都适用。")])

    heading(doc, "3.1 安装 .NET 10 SDK", 2)
    para(doc, [C("开发最小 API，你需要安装 "), B(".NET 10 SDK"), C("。注意区分 SDK 和运行时（Runtime）：")])
    img(doc, "fig_ch3_sdk_runtime.png", "图 3-1  SDK 与运行时的区别")
    bullet(doc, [B("SDK（Software Development Kit）"), C("：开发工具包，包含编译器、CLI 命令行工具，"
                 "还内含运行时。"), B("开发阶段装它。")])
    bullet(doc, [B("Runtime（运行时）"), C("：只负责运行已经编译好的程序。"), B("部署到服务器时装它即可。")])
    para(doc, [C("到微软官网 "), M("https://dotnet.microsoft.com/download"),
               C(" 下载对应操作系统（Windows / macOS / Linux）的 .NET 10 SDK 安装即可。"
                 "装好后，打开终端验证：")])
    code_block(doc, '''dotnet --version''', lang="bash")
    para(doc, [C("如果输出类似 "), M("10.0.100"), C(" 这样的版本号，说明安装成功。再用下面这条命令可以看到已安装的所有 SDK：")])
    code_block(doc, '''dotnet --list-sdks''', lang="bash")
    note(doc, "提示", "Windows 用户如果安装了 Visual Studio 2026，勾选“ASP.NET 和 Web 开发”工作负载时，"
         ".NET 10 SDK 会一并装好，无需单独下载。", "tip")

    heading(doc, "3.2 用 CLI 创建项目 dotnet new web", 2)
    para(doc, [C("在终端里，进入你想放代码的目录，执行下面这条命令，创建一个名为 "),
               M("MyApi"), C(" 的最小 API 项目：")])
    code_block(doc, '''dotnet new web -n MyApi''', lang="bash")
    para(doc, "这条命令拆开看：")
    bullet(doc, [M("dotnet new"), C("：创建新项目。")])
    bullet(doc, [M("web"), C("：使用“空的 ASP.NET Core（最小 API）”模板。这正是我们要的最精简模板。")])
    bullet(doc, [M("-n MyApi"), C("：指定项目名为 MyApi（会创建同名文件夹）。")])
    para(doc, [C("创建完成后，进入项目文件夹：")])
    code_block(doc, '''cd MyApi''', lang="bash")
    note(doc, "小知识", "还有一个常见模板叫 webapi（dotnet new webapi），它会额外生成一个天气预报示例和 OpenAPI 配置。"
         "本教程为了从最简开始，先用最精简的 web 模板。", "info")

    heading(doc, "3.3 项目结构详解", 2)
    para(doc, "用 web 模板创建的项目非常精简，主要就下面这几个文件。先认识它们，后面才知道该改哪里：")
    img(doc, "fig_ch3_structure.png", "图 3-2  最小 API 项目结构")
    bullet(doc, [B("Program.cs"), C("：程序入口，也是你"), B("写接口的主战场"), C("，本教程绝大部分代码都写在这里。")])
    bullet(doc, [B("MyApi.csproj"), C("：项目文件，声明目标框架（如 "), M("net10.0"),
                 C("）和引用的 NuGet 包。")])
    bullet(doc, [B("appsettings.json"), C("：配置文件，存放数据库连接串、自定义参数等（第 14 章详讲）。")])
    bullet(doc, [B("Properties/launchSettings.json"), C("：本地启动配置，比如监听哪个端口、用什么环境。")])
    bullet(doc, [B("bin/ 和 obj/"), C("：编译时自动生成的目录，不用手动管，通常也不提交到 Git。")])
    para(doc, [C("打开 "), M("MyApi.csproj"), C("，内容大致如下：")])
    code_block(doc, '''<Project Sdk="Microsoft.NET.Sdk.Web">
  <PropertyGroup>
    <TargetFramework>net10.0</TargetFramework>
    <Nullable>enable</Nullable>
    <ImplicitUsings>enable</ImplicitUsings>
  </PropertyGroup>
</Project>''', lang="xml")
    bullet(doc, [M("TargetFramework"), C(" 为 "), M("net10.0"), C(" 表示使用 .NET 10。")])
    bullet(doc, [M("ImplicitUsings"), C(" 开启后，常用的 using 会自动引入，所以 Program.cs 顶部看不到一堆 using 也能正常编译。")])

    heading(doc, "3.4 运行与热重载 dotnet watch", 2)
    para(doc, [C("在项目目录下，用下面这条命令运行程序：")])
    code_block(doc, '''dotnet run''', lang="bash")
    para(doc, [C("终端会输出类似这样的日志，告诉你服务器在哪个地址监听：")])
    code_block(doc, '''Now listening on: http://localhost:5000
Application started. Press Ctrl+C to shut down.''', lang="text")
    para(doc, [C("这时打开浏览器访问那个地址即可。要停止服务器，按 "), M("Ctrl + C"), C("。")])
    para(doc, [C("开发时更推荐用"), B("热重载"), C("命令——改完代码保存后自动重新编译运行，不用手动重启：")])
    code_block(doc, '''dotnet watch''', lang="bash")
    img(doc, "fig_ch3_run_flow.png", "图 3-3  从创建到运行的全流程")

    heading(doc, "3.5 Hello World 全流程示例", 2)
    para(doc, "现在把前面所有步骤串起来，完整走一遍。目标：做一个能返回文字和 JSON 的接口。")
    heading(doc, "第 1 步：打开 Program.cs，替换为以下内容", 3)
    code_block(doc, '''var builder = WebApplication.CreateBuilder(args);
var app = builder.Build();

// 接口1：返回一段纯文本
app.MapGet("/", () => "Hello World! 我的第一个最小 API 跑起来了！");

// 接口2：返回 JSON（匿名对象会被自动序列化）
app.MapGet("/api/hello", (string? name) =>
    new { message = $"你好，{name ?? "陌生人"}！", time = DateTime.Now });

app.Run();''')
    heading(doc, "第 2 步：运行", 3)
    code_block(doc, '''dotnet watch''', lang="bash")
    heading(doc, "第 3 步：在浏览器里测试", 3)
    para(doc, "假设服务器监听在 5000 端口，分别访问下面两个地址：")
    bullet(doc, [C("访问 "), M("http://localhost:5000/"), C(" → 页面显示：Hello World! 我的第一个最小 API 跑起来了！")])
    bullet(doc, [C("访问 "), M("http://localhost:5000/api/hello?name=小明"), C(" → 返回一段 JSON：")])
    code_block(doc, '''{
  "message": "你好，小明！",
  "time": "2026-07-17T14:30:00.123456+08:00"
}''', lang="json")
    para(doc, [C("看到 JSON 正常返回，恭喜你——"), B("你已经独立完成了一个真正的最小 API 接口"),
               C("！其中 "), M("(string? name)"),
               C(" 是从查询字符串自动接收的参数，这种“参数绑定”的魔法我们会在第 5 章专门讲。")])
    note(doc, "本章小结", "开发装 SDK、部署装 Runtime；用 dotnet new web 建项目、dotnet watch 热重载运行；"
         "项目里最重要的文件是 Program.cs。你现在已经能创建、运行、测试一个最小 API 了。"
         "从下一章（第 4 章）开始，我们逐个攻克路由、参数、返回值等核心功能。"
         "在继续之前，如果对“前端怎么调用这些接口”还没概念，强烈建议先读一遍书末的【附录 A】。", "key")
    page_break(doc)


# ===========================================================================
def build_appendix_a(doc):
    heading(doc, "附录 A  前端与后端如何协作", 1)
    para(doc, [C("很多新手在学最小 API 时会有一个大大的困惑："),
               B("我写的接口，页面在哪里？用户怎么看到界面？"),
               C("这篇附录就专门解答——把“前端、后端、它们怎么互相调用”彻底讲清楚，"
                 "并"), B("汇总前端调用后端的所有常见方法，每种都给一个例子"),
               C("，方便你按需查阅。")])

    heading(doc, "A.1 前端、后端分别是什么", 2)
    para(doc, [C("一个完整的网络应用，通常分成"), B("前端（Frontend）"), C("和"),
               B("后端（Backend）"), C("两大部分，各司其职、运行在不同的地方：")])
    img(doc, "fig_2_1_front_back_roles.png", "图 A-1  前端与后端的职责与运行位置")
    table(doc, ["", "前端", "后端（最小 API）"],
          [["职责", "显示界面、收集用户操作", "处理数据、业务逻辑、存数据库"],
           ["技术", "HTML/CSS/JS、Vue、React…", "C# + ASP.NET Core"],
           ["运行位置", "用户的浏览器 / 手机", "服务器"],
           ["产出", "用户看到的页面", "通常是 JSON 数据"]])

    heading(doc, "A.2 最小 API 的定位：只提供数据，不管界面", 2)
    para(doc, [B("这是最关键的一句话："), C("最小 API 属于"), B("后端"),
               C("，它"), B("不返回网页"), C("，而是返回"), B("数据（通常是 JSON）"),
               C("。前端拿到数据后，自己决定怎么显示成界面。例如前端要“商品列表”，接口返回的是这样的数据：")])
    code_block(doc, '''[
  { "id": 1, "name": "苹果", "price": 5.5 },
  { "id": 2, "name": "香蕉", "price": 3.0 }
]''', lang="json")
    note(doc, "什么是 JSON", "JSON 是一种轻量的数据格式，用“键:值”方式组织数据，几乎所有编程语言都能读写，"
         "因此成为前后端之间传递数据的通用语言。", "info")

    heading(doc, "A.3 HTTP 协议入门", 2)
    para(doc, [C("前端和后端是两个独立程序，靠"), B("HTTP 协议"), C("通信。一次“对话”包含几个关键要素：")])
    heading(doc, "① 请求方法（动词）", 3)
    table(doc, ["方法", "含义", "对应操作", "举例"],
          [["GET", "获取数据", "查", "获取商品列表"],
           ["POST", "新增数据", "增", "新建一个商品"],
           ["PUT", "更新数据", "改", "修改某个商品"],
           ["DELETE", "删除数据", "删", "删除某个商品"]])
    heading(doc, "② URL 地址", 3)
    code_block(doc, '''https://api.shop.com/api/products/1
└─┬─┘   └────┬─────┘└─────┬──────┘
 协议       服务器地址        路径（定位到具体接口和资源）''', lang="text")
    heading(doc, "③ 请求头 / 请求体 / 状态码", 3)
    bullet(doc, [B("请求头（Headers）"), C("：附带说明信息，如数据格式、身份令牌。")])
    bullet(doc, [B("请求体（Body）"), C("：新增/修改时提交的数据（通常是 JSON）。")])
    bullet(doc, [B("状态码（Status Code）"), C("：后端返回的“结果代号”：")])
    table(doc, ["状态码", "含义"],
          [["200 OK", "成功"], ["201 Created", "创建成功"],
           ["400 Bad Request", "请求有误"], ["401 Unauthorized", "未登录 / 未授权"],
           ["404 Not Found", "找不到资源"], ["500 Internal Server Error", "服务器内部出错"]])

    heading(doc, "A.4 一次完整调用的全流程", 2)
    para(doc, "把上面的要素串起来，一次完整调用就像“点外卖”：你（前端）下单，餐厅（后端）接单做菜送回，你收到后开吃。")
    img(doc, "fig_2_2_http_flow.png", "图 A-2  一次完整调用的全流程")

    heading(doc, "A.5 前后端分离 vs 前后端一体", 2)
    para(doc, [C("目前最主流的是"), B("前后端分离"),
               C("：前后端是两个独立项目，甚至由不同的人、不同语言开发，只靠约定好的接口沟通。最小 API 天生适合做这种架构里的后端。")])
    img(doc, "fig_2_3_separation.png", "图 A-3  前后端分离架构")

    heading(doc, "A.6 前端调用后端的各种方法大全（各举一例）", 2)
    para(doc, [C("“前端调用后端”本质上都是"), B("发一个 HTTP 请求"),
               C("，只是不同技术栈有不同的写法。下面把常见的方法分门别类，每种都给一个调用同一个接口"),
               M("GET /api/products"), C(" 的最小示例。")])
    img(doc, "fig_appx_methods.png", "图 A-4  前端调用后端 API 的各种方法")

    heading(doc, "方法 1：原生 fetch（现代浏览器首选）", 3)
    para(doc, [M("fetch"), C(" 是浏览器内置的函数，无需引入任何库，是目前最推荐的方式。")])
    code_block(doc, '''// GET 请求
const res = await fetch("https://localhost:5001/api/products");
const data = await res.json();   // 解析 JSON
console.log(data);

// POST 请求（提交数据）
await fetch("https://localhost:5001/api/products", {
  method: "POST",
  headers: { "Content-Type": "application/json" },
  body: JSON.stringify({ id: 3, name: "西瓜", price: 12.0 })
});''', lang="javascript")

    heading(doc, "方法 2：XMLHttpRequest（XHR，传统写法）", 3)
    para(doc, [C("fetch 出现之前的老方式，代码较啰嗦，现在了解即可。很多老项目仍在用。")])
    code_block(doc, '''const xhr = new XMLHttpRequest();
xhr.open("GET", "https://localhost:5001/api/products");
xhr.onload = function () {
  if (xhr.status === 200) {
    const data = JSON.parse(xhr.responseText);
    console.log(data);
  }
};
xhr.send();''', lang="javascript")

    heading(doc, "方法 3：axios（最流行的第三方库）", 3)
    para(doc, [C("axios 是一个广受欢迎的 HTTP 库，API 简洁、自动转 JSON、错误处理方便，Vue/React 项目里很常见。需先安装："),
               M("npm install axios"), C("。")])
    code_block(doc, '''import axios from "axios";

// GET
const res = await axios.get("https://localhost:5001/api/products");
console.log(res.data);   // axios 自动解析好了 JSON

// POST
await axios.post("https://localhost:5001/api/products",
  { id: 3, name: "西瓜", price: 12.0 });''', lang="javascript")

    heading(doc, "方法 4：jQuery.ajax（老项目常见）", 3)
    para(doc, [C("如果项目里还在用 jQuery，通常用 "), M("$.ajax"), C(" 或简写的 "),
               M("$.get / $.post"), C(" 调用接口。")])
    code_block(doc, '''$.ajax({
  url: "https://localhost:5001/api/products",
  method: "GET",
  success: function (data) { console.log(data); }
});

// 简写
$.get("https://localhost:5001/api/products", function (data) {
  console.log(data);
});''', lang="javascript")

    heading(doc, "方法 5：原生 HTML 表单（<form> 提交）", 3)
    para(doc, [C("最传统的方式，不写 JavaScript，靠浏览器提交表单。适合简单场景，但会导致"),
               B("整页刷新/跳转"), C("，体验不如上面的方式。")])
    code_block(doc, '''<form action="https://localhost:5001/api/products" method="POST">
  <input name="name" value="西瓜" />
  <input name="price" value="12.0" />
  <button type="submit">提交</button>
</form>''', lang="html")

    heading(doc, "方法 6：Vue 框架中调用", 3)
    para(doc, [C("Vue 本身不规定用哪个请求库，通常在组件里用 "), M("fetch"), C(" 或 "),
               M("axios"), C("。下面用 Vue 3 组合式 API + fetch：")])
    code_block(doc, '''<script setup>
import { ref, onMounted } from "vue";
const products = ref([]);

onMounted(async () => {
  const res = await fetch("https://localhost:5001/api/products");
  products.value = await res.json();
});
</script>

<template>
  <ul>
    <li v-for="p in products" :key="p.id">{{ p.name }} - ￥{{ p.price }}</li>
  </ul>
</template>''', lang="html")

    heading(doc, "方法 7：React 框架中调用", 3)
    para(doc, [C("React 里常用 "), M("useEffect"), C(" 在组件加载时发请求，用 "),
               M("useState"), C(" 保存数据：")])
    code_block(doc, '''import { useState, useEffect } from "react";

function ProductList() {
  const [products, setProducts] = useState([]);

  useEffect(() => {
    fetch("https://localhost:5001/api/products")
      .then(res => res.json())
      .then(data => setProducts(data));
  }, []);

  return (
    <ul>
      {products.map(p => <li key={p.id}>{p.name} - ￥{p.price}</li>)}
    </ul>
  );
}''', lang="javascript")

    heading(doc, "方法 8：Angular 的 HttpClient", 3)
    para(doc, [C("Angular 内置了 "), M("HttpClient"), C(" 服务，配合 RxJS 的 Observable 使用：")])
    code_block(doc, '''import { HttpClient } from "@angular/common/http";
import { Component, OnInit } from "@angular/core";

@Component({ selector: "app-products", template: `
  <li *ngFor="let p of products">{{ p.name }} - ￥{{ p.price }}</li>` })
export class ProductsComponent implements OnInit {
  products: any[] = [];
  constructor(private http: HttpClient) {}

  ngOnInit() {
    this.http.get<any[]>("https://localhost:5001/api/products")
      .subscribe(data => this.products = data);
  }
}''', lang="typescript")

    heading(doc, "方法 9：微信小程序 wx.request", 3)
    para(doc, [C("微信小程序不能用 fetch，要用它自己的 "), M("wx.request"), C(" API：")])
    code_block(doc, '''wx.request({
  url: "https://localhost:5001/api/products",
  method: "GET",
  success(res) {
    console.log(res.data);   // 后端返回的 JSON
  }
});''', lang="javascript")

    heading(doc, "方法 10：EventSource（接收 SSE 实时推送）", 3)
    para(doc, [C("当后端用 Server-Sent Events（第 13 章）持续推送数据时，前端用浏览器内置的 "),
               M("EventSource"), C(" 来接收，实现“服务器主动推、前端实时收”：")])
    code_block(doc, '''const source = new EventSource("https://localhost:5001/api/stream");
source.onmessage = function (event) {
  console.log("收到推送：", event.data);
};''', lang="javascript")

    heading(doc, "方法 11：WebSocket（双向实时通信）", 3)
    para(doc, [C("需要"), B("双向"), C("实时通信（如聊天、协同）时用 WebSocket，前后端可随时互相发消息：")])
    code_block(doc, '''const ws = new WebSocket("wss://localhost:5001/ws");
ws.onopen = () => ws.send("你好，服务器");
ws.onmessage = (event) => console.log("收到：", event.data);''', lang="javascript")

    heading(doc, "方法 12：服务端调用（Node.js / C# 等）", 3)
    para(doc, [C("“前端”不一定是浏览器——另一个后端服务也可能调用你的接口（服务间调用）。例如 Node.js 用 fetch，"
                 "C# 用 "), M("HttpClient"), C("：")])
    code_block(doc, '''// Node.js（18+ 内置 fetch）
const res = await fetch("https://localhost:5001/api/products");
const data = await res.json();

// C# 服务端 / Blazor
using var http = new HttpClient();
var products = await http.GetFromJsonAsync<Product[]>(
    "https://localhost:5001/api/products");''', lang="javascript")

    para(doc, "把这些方法横向对比一下，方便你按场景选择：")
    table(doc, ["方法", "适用场景", "是否需引入库"],
          [["fetch", "现代浏览器，通用首选", "否（内置）"],
           ["XMLHttpRequest", "维护老项目", "否（内置）"],
           ["axios", "需要更好用的 API/拦截器", "是"],
           ["jQuery.ajax", "已用 jQuery 的老项目", "是（jQuery）"],
           ["HTML 表单", "极简单、可接受整页跳转", "否"],
           ["Vue / React / Angular", "对应框架的前端项目", "框架自带或配 axios"],
           ["wx.request", "微信小程序", "否（小程序内置）"],
           ["EventSource", "接收服务器单向实时推送", "否（内置）"],
           ["WebSocket", "双向实时通信", "否（内置）"],
           ["HttpClient / fetch (服务端)", "服务间调用", "视语言而定"]])

    heading(doc, "A.7 没有前端时如何测试接口", 2)
    para(doc, "学习阶段可以先只写后端。验证接口是否正确，有很多趁手的工具：")
    img(doc, "fig_2_4_test_tools.png", "图 A-5  测试接口的常用工具")
    bullet(doc, [B("浏览器地址栏"), C("：直接输地址即可测，但只能测 GET。")])
    bullet(doc, [B("Postman"), C("：最常用的图形化接口测试工具，各种方法、参数都能方便构造。")])
    bullet(doc, [B(".http 文件"), C("：在 Visual Studio 2026 / VS Code 里新建 "), M(".http"), C(" 文件，写几行就能发请求：")])
    code_block(doc, '''GET https://localhost:5001/api/products

###

POST https://localhost:5001/api/products
Content-Type: application/json

{ "id": 3, "name": "西瓜", "price": 12.0 }''', lang="http")
    bullet(doc, [B("Swagger UI"), C("：.NET 10 内置 OpenAPI 支持，能自动生成可视化测试页面（第 10 章详讲）。")])

    heading(doc, "A.8 跨域（CORS）初识", 2)
    para(doc, [C("浏览器有一条安全规则："), B("从 A 网址打开的页面，默认不允许向 B 网址发请求"),
               C("。比如前端在 "), M("http://localhost:3000"), C("，后端 API 在 "),
               M("https://localhost:5001"), C("，地址不同就会被浏览器拦截。")])
    para(doc, [C("解决办法是后端主动“开门”，声明允许哪些来源访问，这套机制叫 "),
               B("CORS（跨域资源共享）"), C("。学习阶段可以先这样全部放开：")])
    code_block(doc, '''var builder = WebApplication.CreateBuilder(args);
builder.Services.AddCors(options =>
    options.AddDefaultPolicy(p =>
        p.AllowAnyOrigin().AllowAnyMethod().AllowAnyHeader()));

var app = builder.Build();
app.UseCors();   // 启用 CORS''')
    note(doc, "注意", "AllowAnyOrigin() 只适合本地学习。正式上线时应只允许你自己前端的确切地址，"
         "否则有安全风险。CORS 的完整配置见第 15 章。", "warn")
    note(doc, "附录小结", "记住三句话：①最小 API 是后端，只提供数据（JSON），不管界面；"
         "②前端和后端是两个独立程序，通过 HTTP 请求沟通；③无论 fetch、axios、各类框架还是小程序，"
         "“调用”的本质都是发一个 HTTP 请求、拿回 JSON、再渲染。理解了这套协作方式，"
         "你就能把本教程学到的后端接口，接到任何前端上去了。", "key")


# ===========================================================================
def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
    sec.left_margin = sec.right_margin = Inches(0.9)
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_EN
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)

    add_footer(doc)
    build_cover(doc)
    build_toc(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_appendix_a(doc)

    out = "/projects/sandbox/tutorial/ASP.NET-Core-最小API完全教程.docx"
    doc.save(out)
    print("SAVED:", out)


if __name__ == "__main__":
    main()
