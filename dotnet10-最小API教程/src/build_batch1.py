# -*- coding: utf-8 -*-
"""生成 Word 文档：封面 + 目录 + 第1章 + 第2章"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import docbuilder as db
from docbuilder import (heading, para, bullet, numbered, code_block, note,
                        image, table, page_break, inline_code, _set_run,
                        BODY_CN, BODY_EN, CLR_TITLE, CLR_H1, CLR_CAPTION, IMG_DIR)

C = lambda t: (t, {})
B = lambda t: (t, {"bold": True})
M = lambda t: (t, {"mono": True})


def img(doc, name, caption, width=6.0):
    image(doc, os.path.join(IMG_DIR, name), caption, width)


def add_footer(doc):
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《ASP.NET Core 最小 API 完全教程（.NET 10）》")
    _set_run(r, BODY_EN, BODY_CN, 8.5, color=CLR_CAPTION)


# ===========================================================================
def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ASP.NET Core")
    _set_run(r, BODY_EN, BODY_CN, 30, bold=True, color=CLR_TITLE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("最小 API 完全教程")
    _set_run(r, BODY_EN, BODY_CN, 34, bold=True, color=CLR_TITLE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Minimal API · 基于 .NET 10")
    _set_run(r, BODY_EN, BODY_CN, 18, color=db.CLR_H2)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("从零基础到进阶实战 · 图文并茂 · 大量可运行示例")
    _set_run(r, BODY_EN, BODY_CN, 13, italic=True, color=db.CLR_BODY)

    for _ in range(6):
        doc.add_paragraph()
    for line in ["面向读者：完全没接触过最小 API 的新手，以及想系统梳理的开发者",
                 "环境版本：.NET 10 SDK / C# 14",
                 "编写工具：Visual Studio 2026 或 VS Code 均可"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        _set_run(r, BODY_EN, BODY_CN, 11, color=db.CLR_CAPTION)
    page_break(doc)


def build_toc(doc):
    heading(doc, "目录", 1)
    toc = [
        ("第一部分  入门与原理", None),
        ("第 1 章  认识最小 API", ["1.1 什么是最小 API", "1.2 最小 API vs 传统 MVC 控制器",
            "1.3 为什么 .NET 10 推荐用最小 API", "1.4 .NET 10 带来的新变化", "1.5 适用场景与不适用场景"]),
        ("第 2 章  前端与后端如何协作", ["2.1 前端、后端分别是什么", "2.2 最小 API 的定位：只提供数据",
            "2.3 前后端沟通的语言：HTTP 协议入门", "2.4 一次完整调用的全流程", "2.5 实战：前端用 fetch 调用最小 API",
            "2.6 前后端分离 vs 前后端一体", "2.7 没有前端时如何测试接口", "2.8 跨域（CORS）问题初识"]),
        ("第 3 章  运行原理与请求处理管线", None),
        ("第 4 章  环境搭建与第一个程序", None),
        ("第二部分  核心功能", None),
        ("第 5 章  路由（Routing）", None),
        ("第 6 章  参数绑定（Parameter Binding）", None),
        ("第 7 章  返回结果（Results）", None),
        ("第 8 章  依赖注入（DI）", None),
        ("第三部分  数据与校验", None),
        ("第 9 章  请求校验（.NET 10 新特性重点）", None),
        ("第 10 章  数据持久化与 EF Core", None),
        ("第四部分  文档、安全与实时通信", None),
        ("第 11 章  OpenAPI 与接口文档（.NET 10 新特性）", None),
        ("第 12 章  认证与授权", None),
        ("第 13 章  端点过滤器与中间件", None),
        ("第 14 章  实时通信：Server-Sent Events（.NET 10 新特性）", None),
        ("第五部分  工程化与进阶", None),
        ("第 15 章  配置、日志与错误处理", None),
        ("第 16 章  跨域（CORS）、限流与健康检查", None),
        ("第 17 章  API 版本控制", None),
        ("第 18 章  测试", None),
        ("第 19 章  组织大型项目", None),
        ("第 20 章  部署与性能", None),
        ("第 21 章  综合实战：从零构建一个完整 API", None),
        ("附录 A  最小 API 常用速查表", None),
        ("附录 B  从 Controller 迁移到最小 API 对照表", None),
        ("附录 C  常见错误与排查", None),
    ]
    for title, subs in toc:
        is_part = title.startswith("第") and "部分" in title
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2 if subs else 3)
        r = p.add_run(title)
        if is_part:
            p.paragraph_format.space_before = Pt(10)
            _set_run(r, BODY_EN, BODY_CN, 13, bold=True, color=CLR_TITLE)
        else:
            _set_run(r, BODY_EN, BODY_CN, 11.5, bold=True, color=CLR_H1)
        if subs:
            for s in subs:
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Pt(22)
                sp.paragraph_format.space_after = Pt(1)
                sr = sp.add_run(s)
                _set_run(sr, BODY_EN, BODY_CN, 10, color=db.CLR_BODY)
    note(doc, "阅读提示", "本教程前四章重在讲清“是什么、为什么、怎么跑起来”，"
         "从第 5 章开始进入具体功能。每章都配有示意图和可直接复制运行的代码。"
         "标有“.NET 10 新特性”的章节，是相较旧版本的重要升级点，建议重点掌握。", "info")
    page_break(doc)


# ===========================================================================
def build_chapter1(doc):
    heading(doc, "第 1 章  认识最小 API", 1)
    para(doc, [C("在写第一行代码之前，我们先花一章时间把两个最基本的问题弄明白："),
               B("最小 API 到底是什么？"), C("以及"), B("我为什么要用它？"),
               C("想清楚这两点，后面学起来才不会“知其然不知其所以然”。")])

    # 1.1
    heading(doc, "1.1 什么是最小 API（Minimal API）", 2)
    para(doc, [C("“最小 API”是 ASP.NET Core 提供的一种"), B("用最少的代码构建 HTTP 接口"),
               C("的方式。它的英文是 "), M("Minimal API"),
               C("，从 .NET 6 引入，到 .NET 10 已经非常成熟，成为微软官方推荐的 API 构建方式。")])
    para(doc, [C("“最小”体现在：你不需要建一堆文件、写一堆类，"),
               B("一个 Program.cs 文件、几行代码"), C("，就能跑起来一个真正的 Web 接口。看下面这段完整程序：")])
    code_block(doc, '''var builder = WebApplication.CreateBuilder(args);
var app = builder.Build();

// 定义一个接口：访问 /hello 时返回一段文字
app.MapGet("/hello", () => "你好，最小 API！");

app.Run();  // 启动服务器，开始监听请求''')
    para(doc, [C("就这么几行，运行后在浏览器打开 "), M("http://localhost:5000/hello"),
               C("，你就能看到 “你好，最小 API！”。这就是一个完整的后端接口了——"),
               B("没有 Controller 类，没有一堆配置文件，没有繁琐的约定。")])
    note(doc, "名词解释", "“API”指应用程序编程接口，这里可以简单理解为“一个可以通过网址访问、"
         "用来交换数据的服务入口”。“HTTP API”就是通过 HTTP 协议访问的接口。", "info")

    # 1.2
    heading(doc, "1.2 最小 API vs 传统 MVC 控制器", 2)
    para(doc, [C("在最小 API 出现之前，ASP.NET Core 构建接口的主流方式是"),
               B("MVC 控制器（Controller）"), C("。两者都能做出一样的接口，区别在于“写法的繁简”和“适用的规模”。")])
    img(doc, "fig_1_1_minimal_vs_mvc.png", "图 1-1  最小 API 与传统 MVC 控制器的对比")
    para(doc, [C("下面用一个“返回商品列表”的例子直观对比。先看"), B("传统 MVC 控制器"), C("的写法：")])
    code_block(doc, '''// 传统 MVC：需要单独的 Controller 类、特性标注、约定命名
[ApiController]
[Route("api/[controller]")]
public class ProductsController : ControllerBase
{
    [HttpGet]
    public IActionResult GetAll()
    {
        var products = new[] { "苹果", "香蕉" };
        return Ok(products);
    }
}''')
    para(doc, [C("再看"), B("最小 API"), C("实现同样的功能：")])
    code_block(doc, '''// 最小 API：一行搞定，无需类、无需特性
app.MapGet("/api/products", () => new[] { "苹果", "香蕉" });''')
    para(doc, [C("两者对外暴露的接口效果"), B("完全一样"),
               C("，但最小 API 明显更简洁。下表总结它们的核心差异：")])
    table(doc,
          ["对比项", "最小 API", "MVC 控制器"],
          [["代码量", "极少，几行即可", "较多，需类+特性+约定"],
           ["文件组织", "可集中在 Program.cs", "通常每个控制器一个文件"],
           ["学习曲线", "平缓，直观", "较陡，需理解约定"],
           ["启动/内存", "更快、更省", "相对更重"],
           ["适合规模", "中小型、微服务", "大型、复杂系统"],
           ["功能完整度", "已覆盖绝大多数场景", "最全面"]])
    note(doc, "重点", "最小 API 并不是“功能被阉割的 MVC”。到了 .NET 10，它已经补齐了校验、"
         "OpenAPI 文档、认证授权等企业级能力，绝大多数项目都可以放心使用。", "key")

    # 1.3
    heading(doc, "1.3 为什么 .NET 10 推荐用最小 API", 2)
    para(doc, "微软在 .NET 10 官方文档中，明确把最小 API 列为构建 HTTP 接口的推荐方式。原因主要有四点：")
    numbered(doc, [B("上手快、心智负担低"), C("——新手不用先理解一堆 MVC 约定，看到路由就能对应到代码。")])
    numbered(doc, [B("性能好"), C("——更少的抽象层意味着更快的启动速度和更低的内存占用，天然适合云原生、容器与微服务。")])
    numbered(doc, [B("与现代实践契合"), C("——前后端分离、微服务架构下，后端往往只需提供 JSON 接口，最小 API 正好胜任。")])
    numbered(doc, [B("功能已足够成熟"), C("——经过 .NET 6/7/8/9 的持续增强，到 .NET 10 已是可用于生产的完整框架。")])

    # 1.4
    heading(doc, "1.4 .NET 10 带来的新变化", 2)
    para(doc, "如果你之前用过旧版本的最小 API，或者看过一些旧教程，要特别注意 .NET 10 带来的几项重要升级：")
    bullet(doc, [B("内置模型校验"), C("：不用再自己写校验代码，直接用 "), M("[Required]"),
                 C("、"), M("[Range]"), C(" 等特性即可，且兼容原生 AoT 编译（第 9 章详讲）。")])
    bullet(doc, [B("OpenAPI 3.1 支持"), C("：内置生成符合 OpenAPI 3.1 规范的接口文档，能配合 Swagger UI / Scalar 使用（第 11 章）。")])
    bullet(doc, [B("Server-Sent Events（SSE）"), C("：通过 "), M("TypedResults.ServerSentEvents"),
                 C(" 轻松实现服务器实时推送（第 14 章）。")])
    bullet(doc, [B("原生 AoT 编译增强"), C("：可编译成独立的原生可执行文件，启动更快、体积更小（第 20 章）。")])
    note(doc, "注意", "由于最小 API 每个版本都在快速演进，网上不少旧教程的写法在 .NET 10 里可能已有更简单的替代方案。"
         "本教程所有示例均以 .NET 10 为准。", "warn")

    # 1.5
    heading(doc, "1.5 适用场景与不适用场景", 2)
    para(doc, "任何技术都有它擅长和不擅长的地方。什么时候该用最小 API，什么时候该考虑别的方案？")
    img(doc, "fig_1_2_when_to_use.png", "图 1-2  最小 API 的适用与不适用场景")
    para(doc, [C("简单来说："), B("绝大多数“提供数据接口”的后端项目，最小 API 都是很好的选择"),
               C("。只有在需要服务器端渲染网页、或深度依赖 MVC 特有管线的少数场景，才需要另作考虑。"
                 "对新手而言，先用最小 API 把基础打牢，是最高效的路径。")])
    note(doc, "本章小结", "最小 API 是一种用极少代码构建 HTTP 接口的方式；它比 MVC 更轻更快，"
         "在 .NET 10 中功能已相当完整，是官方推荐的首选。记住一句话：最小 API 负责“提供数据接口”，"
         "而不是“渲染网页”。下一章我们就来看看，它提供的数据，前端是怎么拿去用的。", "key")
    page_break(doc)


# ===========================================================================
def build_chapter2(doc):
    heading(doc, "第 2 章  前端与后端如何协作", 1)
    para(doc, [C("很多新手在学最小 API 时会有一个大大的困惑："),
               B("我写的接口，页面在哪里？用户怎么看到界面？"),
               C("这一章就专门解答这个问题——把“前端、后端、它们怎么互相调用”彻底讲清楚。"
                 "理解了整体协作方式，后面学具体语法就不会迷糊。")])

    # 2.1
    heading(doc, "2.1 前端、后端分别是什么", 2)
    para(doc, [C("一个完整的网络应用，通常分成"), B("前端（Frontend）"), C("和"),
               B("后端（Backend）"), C("两大部分。它们各司其职、运行在不同的地方：")])
    img(doc, "fig_2_1_front_back_roles.png", "图 2-1  前端与后端的职责与运行位置")
    table(doc, ["", "前端", "后端（最小 API）"],
          [["职责", "显示界面、收集用户操作", "处理数据、业务逻辑、存数据库"],
           ["技术", "HTML/CSS/JS、Vue、React…", "C# + ASP.NET Core"],
           ["运行位置", "用户的浏览器 / 手机", "服务器"],
           ["产出", "用户看到的页面", "通常是 JSON 数据"]])

    # 2.2
    heading(doc, "2.2 最小 API 的定位：只提供数据，不管界面", 2)
    para(doc, [B("这是全章最关键的一句话："), C("最小 API 属于"), B("后端"),
               C("，它"), B("不返回网页"), C("，而是返回"), B("数据（通常是 JSON 格式）"),
               C("。前端拿到数据后，自己决定怎么把它显示成漂亮的界面。")])
    para(doc, [C("举个例子，前端向接口要“商品列表”，最小 API 返回的不是一个页面，而是这样一段数据：")])
    code_block(doc, '''[
  { "id": 1, "name": "苹果", "price": 5.5 },
  { "id": 2, "name": "香蕉", "price": 3.0 }
]''', lang="json")
    para(doc, [C("前端拿到这段 JSON 后，可能把它渲染成一个商品卡片列表，也可能画成一张表格——"),
               B("同一份数据，不同前端可以有完全不同的展现方式"), C("。这正是前后端分工的意义所在。")])
    note(doc, "什么是 JSON", "JSON 是一种轻量的数据格式，用“键:值”的方式组织数据，几乎所有编程语言都能读写它，"
         "因此成为前后端之间传递数据的通用语言。", "info")

    # 2.3
    heading(doc, "2.3 前后端沟通的语言：HTTP 协议入门", 2)
    para(doc, [C("前端和后端是两个独立的程序，它们之间靠"), B("HTTP 协议"),
               C("来通信。你可以把 HTTP 理解成它们之间“约定好的对话规则”。一次对话包含几个关键要素：")])
    heading(doc, "① 请求方法（动词）——想干什么", 3)
    para(doc, "HTTP 用不同的“动词”表达不同的意图，最常用的四个正好对应数据的增删改查：")
    table(doc, ["方法", "含义", "对应操作", "举例"],
          [["GET", "获取数据", "查", "获取商品列表"],
           ["POST", "新增数据", "增", "新建一个商品"],
           ["PUT", "更新数据", "改", "修改某个商品"],
           ["DELETE", "删除数据", "删", "删除某个商品"]])
    heading(doc, "② URL 地址——找谁", 3)
    para(doc, [C("URL 指明请求发给哪个接口，例如 "),
               M("https://api.shop.com/api/products/1"), C("。它由几部分组成：")])
    code_block(doc, '''https://api.shop.com/api/products/1
└─┬─┘   └────┬─────┘└─────┬──────┘
 协议       服务器地址        路径（定位到具体接口和资源）''', lang="text")
    heading(doc, "③ 请求头 / 请求体 / 状态码", 3)
    bullet(doc, [B("请求头（Headers）"), C("：附带的说明信息，比如数据格式、身份令牌等。")])
    bullet(doc, [B("请求体（Body）"), C("：新增/修改时，把要提交的数据放这里（通常是 JSON）。")])
    bullet(doc, [B("状态码（Status Code）"), C("：后端返回的“结果代号”，告诉前端成功还是失败：")])
    table(doc, ["状态码", "含义"],
          [["200 OK", "成功"], ["201 Created", "创建成功"],
           ["400 Bad Request", "请求有误（比如参数不对）"],
           ["401 Unauthorized", "未登录 / 未授权"],
           ["404 Not Found", "找不到资源"],
           ["500 Internal Server Error", "服务器内部出错"]])

    # 2.4
    heading(doc, "2.4 一次完整调用的全流程", 2)
    para(doc, [C("把上面的要素串起来，一次完整的前后端调用就像“点外卖”：你（前端）下单，"
                 "餐厅（后端）接单、做菜、送回，你收到后开吃。看下图：")])
    img(doc, "fig_2_2_http_flow.png", "图 2-2  一次完整调用的全流程")
    numbered(doc, [B("前端发请求"), C("：浏览器向 "), M("GET /api/products"), C(" 发出请求，意思是“我要商品列表”。")])
    numbered(doc, [B("后端处理"), C("：最小 API 收到请求，去数据库查数据、组织好结果。")])
    numbered(doc, [B("后端返回响应"), C("：把结果以 JSON 形式返回，并附上状态码 200。")])
    numbered(doc, [B("前端渲染"), C("：前端拿到 JSON，把它显示成用户能看懂的界面。")])

    # 2.5
    heading(doc, "2.5 实战：前端用 fetch 调用最小 API", 2)
    para(doc, "光说不练假把式。下面是一个可以真正跑起来的完整例子：后端用最小 API 提供接口，前端用浏览器内置的 fetch 调用它。")
    heading(doc, "第 1 步：后端（Program.cs）", 3)
    code_block(doc, '''var builder = WebApplication.CreateBuilder(args);

// 允许前端跨域访问（2.8 节会解释为什么需要）
builder.Services.AddCors(options =>
    options.AddDefaultPolicy(p =>
        p.AllowAnyOrigin().AllowAnyMethod().AllowAnyHeader()));

var app = builder.Build();
app.UseCors();

var products = new List<Product>
{
    new(1, "苹果", 5.5m),
    new(2, "香蕉", 3.0m)
};

// 查询所有商品：GET /api/products
app.MapGet("/api/products", () => products);

// 新增一个商品：POST /api/products
app.MapPost("/api/products", (Product p) =>
{
    products.Add(p);
    return Results.Created($"/api/products/{p.Id}", p);
});

app.Run();

record Product(int Id, string Name, decimal Price);''')
    heading(doc, "第 2 步：前端（index.html）", 3)
    code_block(doc, '''<!DOCTYPE html>
<html>
<body>
  <h2>商品列表</h2>
  <ul id="list"></ul>
  <button onclick="addProduct()">添加西瓜</button>

  <script>
    const API = "https://localhost:5001/api/products";

    // 调用后端“查询”接口 (GET)
    async function loadProducts() {
      const res = await fetch(API);       // 发请求
      const data = await res.json();      // 把返回的 JSON 转成 JS 对象
      document.getElementById("list").innerHTML =
        data.map(p => `<li>${p.name} - ￥${p.price}</li>`).join("");
    }

    // 调用后端“新增”接口 (POST)
    async function addProduct() {
      await fetch(API, {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: JSON.stringify({ id: 3, name: "西瓜", price: 12.0 })
      });
      loadProducts();   // 新增后刷新列表
    }

    loadProducts();     // 页面打开就加载一次
  </script>
</body>
</html>''', lang="html")
    para(doc, [C("这里的核心就是 "), M("fetch(API)"), C(" 这一句——它就是“前端调用后端”的动作。"
               "前端发出请求，后端的 "), M("MapGet"), C(" 被触发，返回 JSON，前端再把数据显示出来。整个闭环就跑通了。")])
    note(doc, "小贴士", "前端代码里的 async / await 是 JavaScript 处理“等待网络返回”的语法，"
         "因为请求要花时间，用 await 表示“等它返回再继续”。", "tip")

    # 2.6
    heading(doc, "2.6 前后端分离 vs 前后端一体", 2)
    para(doc, [C("上面这种“前端一个项目、后端一个项目、通过 HTTP 通信”的模式，叫做"),
               B("前后端分离"), C("，是目前最主流的架构：")])
    img(doc, "fig_2_3_separation.png", "图 2-3  前后端分离架构")
    bullet(doc, [B("前后端分离"), C("：前后端是两个独立项目，甚至由不同的人、不同语言开发，只靠约定好的接口沟通。"
                 "最小 API 天生适合做这种架构里的后端。")])
    bullet(doc, [B("前后端一体"), C("：由服务器直接生成完整网页返回（如传统的 Razor Pages / MVC 视图）。这不是最小 API 的主场。")])

    # 2.7
    heading(doc, "2.7 没有前端时如何测试接口", 2)
    para(doc, "学习阶段，你完全可以先只写后端、不写前端。那怎么验证接口对不对？有很多趁手的工具：")
    img(doc, "fig_2_4_test_tools.png", "图 2-4  测试接口的常用工具")
    bullet(doc, [B("浏览器地址栏"), C("：直接输入接口地址即可测试，但只能测 GET 请求。")])
    bullet(doc, [B("Postman"), C("：最常用的图形化接口测试工具，各种请求方法、参数都能方便地构造。")])
    bullet(doc, [B(".http 文件"), C("：在 Visual Studio 2026 / VS Code 里新建一个 "), M(".http"),
                 C(" 文件，写几行就能直接发请求，非常方便（后面章节会用到）。")])
    bullet(doc, [B("Swagger UI"), C("：.NET 10 内置 OpenAPI 支持，能自动生成一个可视化的接口测试页面（第 11 章详讲）。")])

    # 2.8
    heading(doc, "2.8 跨域（CORS）问题初识", 2)
    para(doc, [C("你可能注意到，2.5 节的后端代码里有一段 "), M("AddCors"), C(" 配置。这是为了解决"),
               B("跨域"), C("问题。")])
    para(doc, [C("浏览器有一条安全规则："), B("从 A 网址打开的页面，默认不允许向 B 网址发请求"),
               C("。比如前端页面跑在 "), M("http://localhost:3000"), C("，而后端 API 在 "),
               M("https://localhost:5001"), C("，两者地址不同，浏览器就会拦截这个请求。")])
    para(doc, [C("解决办法是：后端主动“开门”，声明“我允许来自某些地址的请求”，这套机制就叫 "),
               B("CORS（跨域资源共享）"), C("。上面代码里的 "), M("AllowAnyOrigin()"),
               C(" 表示“允许任何来源”，方便学习阶段使用。")])
    note(doc, "注意", "AllowAnyOrigin() 只适合本地学习和测试。正式上线时应当只允许你自己前端的确切地址，"
         "否则存在安全风险。CORS 的完整配置会在第 16 章深入讲解。", "warn")
    note(doc, "本章小结", "记住三句话：①最小 API 是后端，只提供数据（JSON），不管界面；"
         "②前端和后端是两个独立程序，通过 HTTP 请求沟通；③“调用”的本质是前端用 fetch/axios 发请求、"
         "后端处理并返回 JSON、前端再渲染。理解了这套协作方式，我们下一章就深入后端内部，"
         "看看一个请求进入最小 API 后到底经历了什么。", "key")
    page_break(doc)


# ===========================================================================
def main():
    doc = Document()
    # 页面设置 A4
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    # 默认样式字体
    normal = doc.styles["Normal"]
    normal.font.name = BODY_EN
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)

    add_footer(doc)
    build_cover(doc)
    build_toc(doc)
    build_chapter1(doc)
    build_chapter2(doc)

    out = "/projects/sandbox/tutorial/ASP.NET-Core-最小API完全教程.docx"
    doc.save(out)
    print("SAVED:", out)


if __name__ == "__main__":
    main()
